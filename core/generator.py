"""
Word 文档生成器 (增强版)
将 Simulink 模型分析结果生成美观的 Word 文档
增强功能描述：功能逻辑、函数概述、关键参数等
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from typing import Dict, List, Optional
from datetime import datetime
import os
import re

from .parser import SimulinkModel, SubSystem, Block
from .analyzer import ModelAnalyzer


class DocGenerator:
    """Word 文档生成器（增强版）"""
    
    def __init__(self, model: SimulinkModel, analyzer: ModelAnalyzer):
        self.model = model
        self.analyzer = analyzer
        self.doc = Document()
        
        # 设置文档默认样式
        self._setup_styles()
        
        # 设置页面边距
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置正文样式
        style = self.doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 设置标题样式
        for i in range(1, 4):
            style_name = f'Heading {i}'
            if style_name in self.doc.styles:
                style = self.doc.styles[style_name]
                style.font.name = 'Microsoft YaHei'
                style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
                style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                
                if i == 1:
                    style.font.size = Pt(16)
                elif i == 2:
                    style.font.size = Pt(14)
                else:
                    style.font.size = Pt(12)
    
    def _set_table_style(self, table):
        """设置表格样式"""
        table.autofit = True
        
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D9E2F3')
                cell._tc.get_or_add_tcPr().append(shading)
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    def generate(self, output_path: str, options: Dict = None):
        """生成 Word 文档"""
        options = options or {}
        
        # 1. 封面
        self._add_cover()
        
        # 2. 目录占位
        self._add_toc_placeholder()
        
        # 3. 模型概述
        self._add_overview()
        
        # 4. 系统架构
        self._add_architecture()
        
        # 5. 子系统详细说明
        self._add_subsystem_details(options)
        
        # 6. 附录
        if options.get('include_data_dict', True):
            self._add_appendix()
        
        self.doc.save(output_path)
    
    def _add_cover(self):
        """添加封面"""
        for _ in range(3):
            self.doc.add_paragraph()
        
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(self.model.name)
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("模型设计文档")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x5B, 0x5B, 0x5B)
        
        for _ in range(4):
            self.doc.add_paragraph()
        
        info_table = self.doc.add_table(rows=6, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        info_data = [
            ('模型名称', self.model.name),
            ('版本', self.model.version or 'N/A'),
            ('作者', self.model.author or 'N/A'),
            ('创建时间', self._format_date(self.model.created)),
            ('最后修改', self._format_date(self.model.modified)),
            ('文档生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
            for paragraph in info_table.rows[i].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        self.doc.add_page_break()
    
    def _format_date(self, date_str: str) -> str:
        """格式化日期"""
        if not date_str:
            return 'N/A'
        try:
            if 'T' in date_str:
                dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
                return dt.strftime('%Y-%m-%d %H:%M')
            return date_str
        except:
            return date_str
    
    def _add_toc_placeholder(self):
        """添加目录占位"""
        self.doc.add_heading('目录', level=1)
        
        note = self.doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = note.add_run('（在 Word 中右键点击此处，选择"更新域"生成目录）')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        
        self.doc.add_page_break()
    
    def _add_overview(self):
        """添加模型概述"""
        self.doc.add_heading('1. 模型概述', level=1)
        
        overview = self.analyzer.get_model_overview()
        
        self.doc.add_heading('1.1 基本信息', level=2)
        
        info_table = self.doc.add_table(rows=1, cols=4)
        self._set_table_style(info_table)
        
        headers = ['属性', '值', '属性', '值']
        for i, h in enumerate(headers):
            info_table.rows[0].cells[i].text = h
        
        data = [
            ('模型名称', overview['name'], '模块总数', str(overview['total_blocks'])),
            ('版本', overview['version'] or 'N/A', '子系统数量', str(overview['total_subsystems'])),
            ('作者', overview['author'] or 'N/A', '最大层级深度', str(overview['max_depth'])),
            ('创建时间', self._format_date(self.model.created), '修改时间', self._format_date(self.model.modified)),
        ]
        
        for row_data in data:
            row = info_table.add_row()
            for i, val in enumerate(row_data):
                row.cells[i].text = val
        
        self.doc.add_heading('1.2 模块类型统计', level=2)
        
        if overview['block_type_stats']:
            type_table = self.doc.add_table(rows=1, cols=3)
            self._set_table_style(type_table)
            
            headers = ['模块类型', '数量', '示例模块']
            for i, h in enumerate(headers):
                type_table.rows[0].cells[i].text = h
            
            for stat in overview['block_type_stats'][:20]:
                row = type_table.add_row()
                row.cells[0].text = stat.block_type
                row.cells[1].text = str(stat.count)
                row.cells[2].text = ', '.join(stat.examples[:3])
        
        self.doc.add_page_break()
    
    def _add_architecture(self):
        """添加系统架构"""
        self.doc.add_heading('2. 系统架构', level=1)
        
        self.doc.add_heading('2.1 层级结构', level=2)
        
        subsystems = self.analyzer.get_all_subsystems()
        
        hierarchy_table = self.doc.add_table(rows=1, cols=4)
        self._set_table_style(hierarchy_table)
        
        headers = ['层级', '子系统名称', '模块数', '端口(I/O)']
        for i, h in enumerate(headers):
            hierarchy_table.rows[0].cells[i].text = h
        
        for level, sub in subsystems[:50]:
            row = hierarchy_table.add_row()
            row.cells[0].text = str(level)
            row.cells[1].text = '  ' * level + sub.name
            row.cells[2].text = str(len(sub.blocks))
            row.cells[3].text = f'{len(sub.inports)}/{len(sub.outports)}'
        
        if len(subsystems) > 50:
            row = hierarchy_table.add_row()
            row.cells[0].text = '...'
            row.cells[1].text = f'(还有 {len(subsystems) - 50} 个子系统)'
        
        self.doc.add_page_break()
    
    def _add_subsystem_details(self, options: Dict):
        """添加子系统详细说明"""
        self.doc.add_heading('3. 子系统详细说明', level=1)
        
        subsystems = self.analyzer.get_all_subsystems()
        
        for i, (level, subsystem) in enumerate(subsystems):
            if i > 0 and i % 10 == 0:
                self.doc.add_page_break()
            
            title_prefix = '  ' * min(level, 2)
            self.doc.add_heading(f'{title_prefix}3.{i+1} {subsystem.name}', level=2)
            
            summary = self.analyzer.get_subsystem_summary(subsystem)
            
            # 1. 功能描述（增强版）
            self.doc.add_heading('功能描述', level=3)
            func_desc = self._generate_detailed_description(subsystem)
            self.doc.add_paragraph(func_desc)
            
            # 2. 函数概述
            self.doc.add_heading('函数概述', level=3)
            func_overview = self._generate_function_overview(subsystem)
            self.doc.add_paragraph(func_overview)
            
            # 3. 关键参数
            key_params = self._extract_key_parameters(subsystem)
            if key_params:
                self.doc.add_heading('关键参数', level=3)
                self._add_parameters_table(key_params)
            
            # 4. 基本信息（表格）
            self.doc.add_heading('基本信息', level=3)
            
            info_table = self.doc.add_table(rows=2, cols=4)
            self._set_table_style(info_table)
            
            info_table.rows[0].cells[0].text = '层级'
            info_table.rows[0].cells[1].text = str(level)
            info_table.rows[0].cells[2].text = '模块数'
            info_table.rows[0].cells[3].text = str(len(subsystem.blocks))
            
            info_table.rows[1].cells[0].text = '输入端口'
            info_table.rows[1].cells[1].text = str(len(subsystem.inports))
            info_table.rows[1].cells[2].text = '输出端口'
            info_table.rows[1].cells[3].text = str(len(subsystem.outports))
            
            # 5. 输入/输出端口表
            if summary['inports']:
                self.doc.add_heading('输入端口', level=3)
                self._add_port_table(summary['inports'])
            
            if summary['outports']:
                self.doc.add_heading('输出端口', level=3)
                self._add_port_table(summary['outports'])
            
            # 6. 内部模块清单
            if options.get('include_block_list', True) and subsystem.blocks:
                self.doc.add_heading('内部模块', level=3)
                self._add_block_table(subsystem)
            
            self.doc.add_paragraph()
    
    def _generate_detailed_description(self, subsystem: SubSystem) -> str:
        """
        生成详细的功能描述
        包括：功能目标、实现原理、关键公式、各模块作用、处理流程
        """
        descriptions = []
        
        # 1. 功能目标和概述
        goal_desc = self._analyze_function_goal(subsystem)
        descriptions.append(goal_desc)
        
        # 2. 实现原理
        principle_desc = self._analyze_implementation_principle(subsystem)
        if principle_desc:
            descriptions.append(principle_desc)
        
        # 3. 关键公式和原理描述（新增）
        formula_desc = self._get_function_formula(subsystem)
        if formula_desc:
            descriptions.append(formula_desc)
        
        # 4. 各模块具体作用
        module_desc = self._analyze_module_roles(subsystem)
        if module_desc:
            descriptions.append(module_desc)
        
        # 5. 处理流程
        flow_desc = self._analyze_processing_flow(subsystem)
        if flow_desc:
            descriptions.append(flow_desc)
        
        if descriptions:
            return '\n\n'.join(descriptions)
        
        return f'该子系统名为"{subsystem.name}"，包含 {len(subsystem.blocks)} 个内部模块。'
    
    def _analyze_function_goal(self, subsystem: SubSystem) -> str:
        """分析功能目标"""
        name_lower = subsystem.name.lower()
        
        # 根据名称和内容推断功能目标
        goals = {
            'foc': '该子系统实现磁场定向控制（FOC），通过坐标变换和电流环控制，实现对电机转矩和磁通的独立控制，提高电机的动态响应和效率。',
            'currentloop': '该子系统实现电流环控制，通过对电流的精确调节，保证电机输出转矩的准确性和响应速度。',
            'svpwm': '该子系统实现空间矢量脉宽调制（SVPWM），将参考电压矢量转换为三相PWM信号，驱动逆变器工作，提高直流电压利用率。',
            'park': '该子系统实现Park变换（坐标变换），将两相静止坐标系（αβ）转换为两相旋转坐标系，使交流量转换为直流量，便于PI控制。',
            'clark': '该子系统实现Clark变换，将三相静止坐标系（abc）转换为两相静止坐标系（αβ），简化后续的坐标变换和控制。',
            'inv_park': '该子系统实现逆Park变换，将两相旋转坐标系转换回两相静止坐标系，用于将控制量转换回实际物理量。',
            'speedloop': '该子系统实现速度环控制，通过调节电机转速，实现速度的精确跟踪和稳定控制。',
            'torque': '该子系统实现转矩控制，通过对转矩的精确调节，保证电机的输出性能。',
            'observer': '该子系统实现观测器功能，通过测量可观测的量来估计不可直接测量的状态变量，如转子位置、磁链等。',
            'filter': '该子系统实现信号滤波功能，滤除信号中的噪声和干扰，提取有效信号成分。',
            'protect': '该子系统实现保护功能，监测系统运行状态，在异常情况下触发保护动作，防止设备损坏。',
            'sensor': '该子系统实现传感器信号处理功能，对传感器采集的原始信号进行处理和转换。',
        }
        
        for key, goal in goals.items():
            if key in name_lower:
                return goal
        
        # 根据模块组成推断
        has_pi = any(b.block_type in ('Gain', 'Sum', 'Integrator') for b in subsystem.blocks)
        has_transform = any('park' in b.name.lower() or 'clark' in b.name.lower() for b in subsystem.blocks)
        has_pwm = any('svpwm' in b.name.lower() or 'pwm' in b.name.lower() for b in subsystem.blocks)
        
        if has_transform and has_pi:
            return '该子系统实现电机控制的核心算法，通过坐标变换和PI控制实现对电机电流的精确控制。'
        elif has_pi:
            return '该子系统实现闭环控制功能，通过反馈调节实现对控制量的精确跟踪。'
        elif has_pwm:
            return '该子系统实现PWM信号生成功能，用于驱动功率器件。'
        
        return f'该子系统"{subsystem.name}"包含 {len(subsystem.blocks)} 个模块，实现特定的信号处理和控制功能。'
    
    def _analyze_implementation_principle(self, subsystem: SubSystem) -> str:
        """分析实现原理"""
        principles = []
        
        # 分析模块组成
        blocks = subsystem.blocks
        block_types = {b.block_type for b in blocks}
        
        # 检测是否包含坐标变换
        transform_blocks = []
        for b in blocks:
            name_lower = b.name.lower()
            if 'park' in name_lower:
                transform_blocks.append(('Park变换', 'αβ→dq'))
            elif 'clark' in name_lower:
                transform_blocks.append(('Clark变换', 'abc→αβ'))
            elif 'invclark' in name_lower or 'inv_clark' in name_lower:
                transform_blocks.append(('逆Clark变换', 'αβ→abc'))
            elif 'inv' in name_lower and 'park' in name_lower:
                transform_blocks.append(('逆Park变换', 'dq→αβ'))
        
        if transform_blocks:
            transforms_str = '、'.join([f"{t[0]}（{t[1]}）" for t in transform_blocks])
            principles.append(f"**坐标变换**：包含{transforms_str}，实现不同坐标系之间的转换。")
        
        # 检测控制器
        gain_blocks = [b for b in blocks if b.block_type == 'Gain']
        sum_blocks = [b for b in blocks if b.block_type == 'Sum']
        
        if gain_blocks and sum_blocks:
            principles.append(f"**闭环控制**：包含 {len(gain_blocks)} 个增益模块和 {len(sum_blocks)} 个求和模块，构成闭环控制系统。")
            
            # 提取关键增益值
            key_gains = []
            for g in gain_blocks[:3]:
                gain_val = g.parameters.get('Gain', '')
                name_lower = g.name.lower()
                if 'kp' in name_lower:
                    key_gains.append(f"Kp={gain_val}")
                elif 'ki' in name_lower:
                    key_gains.append(f"Ki={gain_val}")
                elif gain_val:
                    key_gains.append(f"{g.name}={gain_val}")
            
            if key_gains:
                principles.append(f"**控制器参数**：{', '.join(key_gains[:3])}。")
        
        # 检测PWM/SVPWM
        for b in blocks:
            if 'svpwm' in b.name.lower():
                principles.append("**调制方式**：采用空间矢量脉宽调制（SVPWM），相比传统SPWM可提高直流电压利用率约15%。")
                break
        
        # 检测滤波器
        if 'Delay' in block_types or 'UnitDelay' in block_types:
            principles.append("**信号处理**：包含延迟单元，用于信号缓存和时序同步。")
        
        if principles:
            return '**实现原理**：\n' + '\n'.join(f"- {p}" for p in principles)
        
        return ""
    
    def _analyze_module_roles(self, subsystem: SubSystem) -> str:
        """分析各模块的具体作用"""
        roles = []
        
        for block in subsystem.blocks[:15]:  # 最多显示15个模块
            role = self._get_block_role(block)
            if role:
                roles.append(f"**{block.name}**：{role}")
        
        if roles:
            return '**各模块作用**：\n' + '\n'.join(f"- {r}" for r in roles)
        
        return ""
    
    def _get_block_role(self, block: Block) -> str:
        """获取模块的具体作用"""
        name_lower = block.name.lower()
        block_type = block.block_type
        
        # 根据名称推断作用
        if 'inport' in block_type.lower() or block_type == 'Inport':
            return "输入端口，接收外部信号"
        elif 'outport' in block_type.lower() or block_type == 'Outport':
            return "输出端口，输出处理结果"
        elif 'park' in name_lower:
            if 'inv' in name_lower:
                return "逆Park变换，将dq旋转坐标系转换为αβ静止坐标系"
            return "Park变换，将αβ静止坐标系转换为dq旋转坐标系"
        elif 'clark' in name_lower:
            if 'inv' in name_lower:
                return "逆Clark变换，将αβ静止坐标系转换为abc三相坐标系"
            return "Clark变换，将abc三相坐标系转换为αβ静止坐标系"
        elif 'svpwm' in name_lower:
            return "空间矢量脉宽调制，生成三相PWM驱动信号"
        elif 'currentloop' in name_lower:
            return "电流环控制，实现电流的精确调节"
        elif block_type == 'Gain':
            gain_val = block.parameters.get('Gain', '')
            if 'kp' in name_lower:
                return f"比例增益，Kp={gain_val}，决定系统的响应速度"
            elif 'ki' in name_lower:
                return f"积分增益，Ki={gain_val}，消除稳态误差"
            elif gain_val:
                return f"增益放大，系数={gain_val}"
            return "信号增益放大"
        elif block_type == 'Sum':
            return "求和运算，实现信号的加减组合"
        elif block_type == 'Product':
            return "乘法运算，实现信号的相乘"
        elif block_type == 'TrigonometricFunction':
            func = block.parameters.get('Operator', 'sin')
            return f"三角函数运算（{func}），用于角度计算"
        elif block_type == 'Delay' or block_type == 'UnitDelay':
            return "延迟单元，实现信号的时间延迟"
        elif block_type == 'Constant':
            val = block.parameters.get('Value', '')
            return f"常量定义，值={val}"
        elif block_type == 'Switch':
            return "条件选择，根据条件切换输出"
        elif block_type == 'RelationalOperator':
            return "比较运算，判断信号的大小关系"
        elif block_type == 'BusSelector':
            return "总线选择器，从总线中提取指定信号"
        elif block_type == 'BusCreator':
            return "总线创建器，将多个信号打包成总线"
        elif block_type == 'Mux':
            return "信号复用，将多个信号合并为一个向量"
        elif block_type == 'Demux':
            return "信号解复用，将向量信号分离为多个标量"
        elif block_type == 'DataTypeConversion':
            return "数据类型转换，统一信号的数据类型"
        
        return ""
    
    def _get_function_formula(self, subsystem: SubSystem) -> str:
        """获取关键功能的公式和原理描述"""
        formulas = []
        
        # 检查是否包含坐标变换
        has_park = False
        has_inv_park = False
        has_clark = False
        has_inv_clark = False
        has_svpwm = False
        has_pi = False
        kp_value = ""
        ki_value = ""
        
        for block in subsystem.blocks:
            name_lower = block.name.lower()
            
            if 'park' in name_lower:
                if 'inv' in name_lower:
                    has_inv_park = True
                else:
                    has_park = True
            elif 'clark' in name_lower:
                if 'inv' in name_lower:
                    has_inv_clark = True
                else:
                    has_clark = True
            elif 'svpwm' in name_lower:
                has_svpwm = True
            
            # 检查 PI 参数
            if block.block_type == 'Gain':
                gain_val = block.parameters.get('Gain', '')
                if 'kp' in name_lower and gain_val:
                    kp_value = gain_val
                    has_pi = True
                elif 'ki' in name_lower and gain_val:
                    ki_value = gain_val
                    has_pi = True
        
        # Park 变换公式
        if has_park:
            formulas.append("""**Park变换公式**：
将两相静止坐标系（αβ）转换为两相旋转坐标系（dq）：
```
Id = Iα·cos(θ) + Iβ·sin(θ)
Iq = -Iα·sin(θ) + Iβ·cos(θ)
```
其中θ为转子电角度。Park变换将交流量转换为直流量，便于PI控制器调节。""")
        
        # 逆 Park 变换公式
        if has_inv_park:
            formulas.append("""**逆Park变换公式**：
将两相旋转坐标系（dq）转换回两相静止坐标系（αβ）：
```
Iα = Id·cos(θ) - Iq·sin(θ)
Iβ = Id·sin(θ) + Iq·cos(θ)
```
用于将控制量转换回实际物理量，供SVPWM调制使用。""")
        
        # Clark 变换公式
        if has_clark:
            formulas.append("""**Clark变换公式**：
将三相静止坐标系（abc）转换为两相静止坐标系（αβ）：
```
Iα = (2/3)·(Ia - 0.5·Ib - 0.5·Ic)
Iβ = (2/3)·(√3/2·Ib - √3/2·Ic)
```
或简化形式（假设Ia+Ib+Ic=0）：
```
Iα = Ia
Iβ = (Ia + 2·Ib)/√3
```
Clark变换将三相系统简化为两相系统，便于后续处理。""")
        
        # 逆 Clark 变换公式
        if has_inv_clark:
            formulas.append("""**逆Clark变换公式**：
将两相静止坐标系（αβ）转换回三相静止坐标系（abc）：
```
Ia = Iα
Ib = -Iα/2 + √3/2·Iβ
Ic = -Iα/2 - √3/2·Iβ
```
用于生成三相调制信号。""")
        
        # SVPWM 原理
        if has_svpwm:
            formulas.append("""**SVPWM原理**：
空间矢量脉宽调制（SVPWM）将参考电压矢量分解为8个基本电压矢量的组合：
- 6个有效矢量：V1-V6，幅值均为2Vdc/3
- 2个零矢量：V0、V7

调制公式：
```
Vref = (T1/Ts)·Vx + (T2/Ts)·V(x+1) + (T0/Ts)·V0
```
其中T1、T2为相邻有效矢量作用时间，T0为零矢量时间，Ts为开关周期。

优势：相比SPWM，直流电压利用率提高约15.47%。""")
        
        # PI 控制器公式
        if has_pi:
            params_str = ""
            if kp_value:
                params_str += f"Kp = {kp_value}"
            if ki_value:
                params_str += f", Ki = {ki_value}" if params_str else f"Ki = {ki_value}"
            
            formulas.append(f"""**PI控制器公式**：
比例-积分控制器实现无静差控制：
```
u(t) = Kp·e(t) + Ki·∫e(τ)dτ
```
离散形式：
```
u(k) = Kp·e(k) + Ki·Ts·Σe(i)
```
其中：
- Kp：比例系数，决定系统响应速度
- Ki：积分系数，消除稳态误差
- e(t)：误差信号
- Ts：采样周期
{f"当前参数：{params_str}" if params_str else ""}

作用：实现电流/转速的无静差跟踪控制。""")
        
        if formulas:
            return '\n\n'.join(formulas)
        
        return ""
    
    def _analyze_processing_flow(self, subsystem: SubSystem) -> str:
        """分析处理流程"""
        # 获取信号流分析
        signal_flow = self.analyzer.analyze_signal_flow(subsystem)
        
        if not signal_flow['signal_chains']:
            return ""
        
        flows = []
        for chain in signal_flow['signal_chains'][:3]:
            input_name = chain.get('input', '')
            path = chain.get('path', [])
            output_name = chain.get('output', '')
            
            if path:
                path_str = ' → '.join(path[:8])  # 最多显示8个步骤
                flows.append(f"{input_name} 经过 {path_str}")
        
        if flows:
            return '**处理流程**：\n' + '\n'.join(f"- {f}" for f in flows)
        
        return ""
    
    def _analyze_name_detailed(self, name: str) -> str:
        """基于名称生成详细的功能描述"""
        name_lower = name.lower()
        
        # 常见控制算法关键词
        control_keywords = {
            'foc': '磁场定向控制（FOC）是一种高效的交流电机控制策略，通过坐标变换将三相交流电转换为直流分量进行控制。',
            'svpwm': '空间矢量脉宽调制（SVPWM）用于生成三相逆变器的开关信号，提高直流电压利用率。',
            'pi': '比例-积分控制器（PI）用于消除稳态误差，提高系统响应性能。',
            'pid': '比例-积分-微分控制器（PID）实现精确的控制调节。',
            'currentloop': '电流环控制，实现对电机电流的快速精确控制。',
            'speedloop': '速度环控制，实现对电机转速的调节。',
            'park': 'Park变换（dq变换），将两相静止坐标系转换为两相旋转坐标系。',
            'clark': 'Clark变换，将三相静止坐标系转换为两相静止坐标系。',
            'inv_park': '逆Park变换，将两相旋转坐标系转换回两相静止坐标系。',
            'invclark': '逆Clark变换，将两相静止坐标系转换回三相静止坐标系。',
            'adc': '模数转换模块，将模拟信号转换为数字信号进行处理。',
            'pwm': '脉宽调制模块，生成控制功率器件的开关信号。',
            'svpwm2': '空间矢量脉宽调制模块，生成三相PWM信号。',
            'current': '电流处理模块，实现电流信号的采集、处理和控制。',
            'voltage': '电压处理模块，实现电压信号的采集、处理和控制。',
            'speed': '速度处理模块，实现速度信号的采集和处理。',
            'torque': '转矩处理模块，实现转矩的计算和控制。',
            'flux': '磁链处理模块，实现磁链的估计和控制。',
            'observer': '观测器模块，用于估计不可直接测量的状态变量。',
            'filter': '滤波器模块，用于信号滤波和噪声抑制。',
            'calculator': '计算模块，实现各种数学运算。',
            'control': '控制模块，实现系统的控制逻辑。',
            'protect': '保护模块，实现系统的过流、过压等保护功能。',
            'sensor': '传感器接口模块，处理传感器信号。',
            'encoder': '编码器接口模块，处理位置/速度反馈信号。',
        }
        
        for key, desc in control_keywords.items():
            if key in name_lower:
                return desc
        
        # 通用描述
        words = re.findall('[A-Z][^A-Z]*', name)
        if words:
            return f"该子系统涉及{''.join(words)}相关功能处理。"
        
        return f'该子系统名为"{name}"。'
    
    def _analyze_function_logic(self, subsystem: SubSystem) -> str:
        """分析功能逻辑"""
        blocks = subsystem.blocks
        
        # 统计模块类型
        block_types = {}
        for block in blocks:
            block_types[block.block_type] = block_types.get(block.block_type, 0) + 1
        
        logic_parts = []
        
        # 分析主要运算模块
        if 'Gain' in block_types:
            gains = [b for b in blocks if b.block_type == 'Gain']
            gain_values = []
            for g in gains:
                if 'Gain' in g.parameters:
                    gain_values.append(f"{g.name}={g.parameters['Gain']}")
            if gain_values:
                logic_parts.append(f"增益运算：包含 {block_types['Gain']} 个增益模块（{', '.join(gain_values[:3])}）")
        
        if 'Sum' in block_types:
            logic_parts.append(f"求和运算：包含 {block_types['Sum']} 个求和模块")
        
        if 'Product' in block_types:
            logic_parts.append(f"乘法运算：包含 {block_types['Product']} 个乘法模块")
        
        if 'Logic' in block_types:
            logic_parts.append(f"逻辑运算：包含 {block_types['Logic']} 个逻辑模块")
        
        if 'Switch' in block_types:
            logic_parts.append(f"条件切换：包含 {block_types['Switch']} 个切换开关")
        
        if 'RelationalOperator' in block_types:
            logic_parts.append(f"比较运算：包含 {block_types['RelationalOperator']} 个比较模块")
        
        if 'TrigonometricFunction' in block_types:
            logic_parts.append(f"三角函数：包含 {block_types['TrigonometricFunction']} 个三角函数模块")
        
        if 'MathFunction' in block_types:
            logic_parts.append(f"数学函数：包含 {block_types['MathFunction']} 个数学函数模块")
        
        # 分析数据流
        if 'BusSelector' in block_types:
            logic_parts.append(f"总线信号选择：从总线中提取 {block_types['BusSelector']} 组信号")
        
        if 'BusCreator' in block_types:
            logic_parts.append(f"总线信号创建：将 {block_types['BusCreator']} 组信号打包成总线")
        
        if 'Mux' in block_types:
            logic_parts.append(f"信号复用：将 {block_types['Mux']} 组信号合并")
        
        if 'Demux' in block_types:
            logic_parts.append(f"信号解复用：将信号分离成 {block_types['Demux']} 路")
        
        # 分析存储和延迟
        if 'Delay' in block_types or 'UnitDelay' in block_types:
            total_delay = block_types.get('Delay', 0) + block_types.get('UnitDelay', 0)
            logic_parts.append(f"延迟处理：包含 {total_delay} 个延迟单元")
        
        if 'Memory' in block_types:
            logic_parts.append(f"存储单元：包含 {block_types['Memory']} 个存储模块")
        
        # 分析数据类型转换
        if 'DataTypeConversion' in block_types:
            logic_parts.append(f"数据类型转换：包含 {block_types['DataTypeConversion']} 个转换模块")
        
        if logic_parts:
            return "功能逻辑：\n" + '\n'.join(f"• {p}" for p in logic_parts)
        
        return ""
    
    def _analyze_signal_flow(self, subsystem: SubSystem) -> str:
        """分析信号流"""
        flow_parts = []
        
        # 输入分析
        if subsystem.inports:
            inport_names = [p.name for p in subsystem.inports]
            flow_parts.append(f"输入信号：{', '.join(inport_names)}")
        
        # 输出分析
        if subsystem.outports:
            outport_names = [p.name for p in subsystem.outports]
            flow_parts.append(f"输出信号：{', '.join(outport_names)}")
        
        # 子系统分析
        child_names = [child.name for child in subsystem.children]
        if child_names:
            flow_parts.append(f"处理流程：信号经过 {', '.join(child_names)} 等子模块处理")
        
        if flow_parts:
            return "信号流：\n" + '\n'.join(f"• {p}" for p in flow_parts)
        
        return ""
    
    def _generate_function_overview(self, subsystem: SubSystem) -> str:
        """生成函数概述"""
        overview_parts = []
        
        # 输入输出关系
        if subsystem.inports and subsystem.outports:
            overview_parts.append(f"该子系统接收 {len(subsystem.inports)} 路输入信号，输出 {len(subsystem.outports)} 路信号。")
        elif subsystem.inports:
            overview_parts.append(f"该子系统接收 {len(subsystem.inports)} 路输入信号。")
        elif subsystem.outports:
            overview_parts.append(f"该子系统输出 {len(subsystem.outports)} 路信号。")
        
        # 主要处理步骤
        block_types = {}
        for block in subsystem.blocks:
            block_types[block.block_type] = block_types.get(block.block_type, 0) + 1
        
        processing_steps = []
        if 'SubSystem' in block_types:
            processing_steps.append(f"包含 {block_types['SubSystem']} 个子系统模块")
        if 'Gain' in block_types or 'Sum' in block_types or 'Product' in block_types:
            processing_steps.append("进行基本数学运算")
        if 'TrigonometricFunction' in block_types:
            processing_steps.append("执行三角函数计算")
        if 'LookupTable' in block_types or 'LookupTable2D' in block_types:
            processing_steps.append("进行查表插值")
        if 'TransferFcn' in block_types:
            processing_steps.append("包含传递函数处理")
        if 'StateSpace' in block_types:
            processing_steps.append("包含状态空间模型")
        
        if processing_steps:
            overview_parts.append("处理步骤：" + "、".join(processing_steps) + "。")
        
        if overview_parts:
            return '\n'.join(overview_parts)
        
        return f"该子系统包含 {len(subsystem.blocks)} 个内部模块，用于信号处理。"
    
    def _extract_key_parameters(self, subsystem: SubSystem) -> List[Dict]:
        """提取关键参数"""
        key_params = []
        
        # 重点关注的关键参数
        key_param_names = [
            'Gain', 'Value', 'Constant', 'Amplitude', 'Frequency',
            'Phase', 'Bias', 'SampleTime', 'UpperLimit', 'LowerLimit',
            'Threshold', 'Numerator', 'Denominator', 'Coefficients',
            'Data', 'Table', 'Breakpoints', 'InitialCondition'
        ]
        
        for block in subsystem.blocks:
            block_params = []
            
            for param_name in key_param_names:
                if param_name in block.parameters:
                    value = block.parameters[param_name]
                    # 简化显示
                    if len(str(value)) > 50:
                        value = str(value)[:50] + '...'
                    block_params.append({
                        'name': param_name,
                        'value': value
                    })
            
            # 只保留有重要参数的模块
            if block_params and block.block_type not in ['Inport', 'Outport', 'SubSystem']:
                key_params.append({
                    'block_name': block.name,
                    'block_type': block.block_type,
                    'params': block_params[:3]  # 每个模块最多显示3个参数
                })
        
        return key_params[:10]  # 最多显示10个模块的参数
    
    def _add_parameters_table(self, key_params: List[Dict]):
        """添加参数表格"""
        if not key_params:
            return
        
        table = self.doc.add_table(rows=1, cols=4)
        self._set_table_style(table)
        
        headers = ['模块名称', '模块类型', '参数名', '参数值']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        
        for item in key_params:
            for param in item['params']:
                row = table.add_row()
                row.cells[0].text = item['block_name']
                row.cells[1].text = item['block_type']
                row.cells[2].text = param['name']
                row.cells[3].text = str(param['value'])
    
    def _add_port_table(self, ports: List[Dict]):
        """添加端口表格"""
        if not ports:
            return
        
        table = self.doc.add_table(rows=1, cols=4)
        self._set_table_style(table)
        
        headers = ['端口名称', '端口号', '数据类型', '描述']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        
        for port in ports:
            row = table.add_row()
            row.cells[0].text = port['name']
            row.cells[1].text = str(port['number'])
            row.cells[2].text = port['data_type'] or 'N/A'
            row.cells[3].text = port['description'] or '-'
    
    def _add_block_table(self, subsystem: SubSystem):
        """添加模块清单表格"""
        if not subsystem.blocks:
            return
        
        blocks_by_type = {}
        for block in subsystem.blocks:
            if block.block_type not in blocks_by_type:
                blocks_by_type[block.block_type] = []
            blocks_by_type[block.block_type].append(block)
        
        table = self.doc.add_table(rows=1, cols=4)
        self._set_table_style(table)
        
        headers = ['模块类型', '数量', '示例', '说明']
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        
        type_descriptions = {
            'Inport': '输入端口',
            'Outport': '输出端口',
            'SubSystem': '子系统',
            'Gain': '增益模块',
            'Sum': '求和模块',
            'Product': '乘法模块',
            'Logic': '逻辑运算',
            'Switch': '切换开关',
            'Delay': '延时模块',
            'UnitDelay': '单位延时',
            'BusSelector': '总线选择器',
            'BusCreator': '总线创建器',
            'Constant': '常量',
            'Reference': '引用模块',
            'DataTypeConversion': '数据类型转换',
            'Demux': '解复用器',
            'Mux': '复用器',
            'Memory': '存储器',
            'TriggerPort': '触发端口',
            'Abs': '绝对值',
            'MinMax': '最大最小值',
            'RelationalOperator': '关系运算',
            'S-Function': 'S函数',
            'Terminator': '终端',
            'From': '信号来源',
            'Goto': '信号去向',
            'TrigonometricFunction': '三角函数',
            'MathFunction': '数学函数',
            'LookupTable': '查表模块',
            'TransferFcn': '传递函数',
            'StateSpace': '状态空间',
        }
        
        for block_type, blocks in blocks_by_type.items():
            row = table.add_row()
            row.cells[0].text = block_type
            row.cells[1].text = str(len(blocks))
            row.cells[2].text = ', '.join([b.name for b in blocks[:3]])
            if len(blocks) > 3:
                row.cells[2].text += f' ...等{len(blocks)}个'
            row.cells[3].text = type_descriptions.get(block_type, '-')
    
    def _add_appendix(self):
        """添加附录"""
        self.doc.add_page_break()
        self.doc.add_heading('附录', level=1)
        
        self.doc.add_heading('A. 数据字典', level=2)
        
        all_inports = []
        all_outports = []
        
        for level, subsystem in self.analyzer.get_all_subsystems():
            for port in subsystem.inports:
                all_inports.append({
                    'subsystem': subsystem.name,
                    'level': level,
                    'port': port
                })
            for port in subsystem.outports:
                all_outports.append({
                    'subsystem': subsystem.name,
                    'level': level,
                    'port': port
                })
        
        if all_inports:
            self.doc.add_heading('A.1 输入端口汇总', level=3)
            
            inport_table = self.doc.add_table(rows=1, cols=5)
            self._set_table_style(inport_table)
            
            headers = ['所属子系统', '端口名称', '端口号', '数据类型', '描述']
            for i, h in enumerate(headers):
                inport_table.rows[0].cells[i].text = h
            
            for item in all_inports[:100]:
                row = inport_table.add_row()
                row.cells[0].text = item['subsystem']
                row.cells[1].text = item['port'].name
                row.cells[2].text = str(item['port'].port_number)
                row.cells[3].text = item['port'].data_type or 'N/A'
                row.cells[4].text = item['port'].description or '-'
            
            if len(all_inports) > 100:
                row = inport_table.add_row()
                row.cells[0].text = '...'
                row.cells[1].text = f'(还有 {len(all_inports) - 100} 个端口)'
        
        if all_outports:
            self.doc.add_heading('A.2 输出端口汇总', level=3)
            
            outport_table = self.doc.add_table(rows=1, cols=5)
            self._set_table_style(outport_table)
            
            headers = ['所属子系统', '端口名称', '端口号', '数据类型', '描述']
            for i, h in enumerate(headers):
                outport_table.rows[0].cells[i].text = h
            
            for item in all_outports[:100]:
                row = outport_table.add_row()
                row.cells[0].text = item['subsystem']
                row.cells[1].text = item['port'].name
                row.cells[2].text = str(item['port'].port_number)
                row.cells[3].text = item['port'].data_type or 'N/A'
                row.cells[4].text = item['port'].description or '-'


def generate_document(model: SimulinkModel, output_path: str, options: Dict = None):
    """便捷函数：生成文档"""
    analyzer = ModelAnalyzer(model)
    generator = DocGenerator(model, analyzer)
    generator.generate(output_path, options)